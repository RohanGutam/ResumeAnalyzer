import streamlit as st
import os
import google.generativeai as genai
import pandas as pd
import docx
from docx import Document
from docx.shared import Pt, RGBColor
import pdfplumber
import spacy
import io
import re
import time
from datetime import datetime
from dotenv import load_dotenv

# ✅ Ensure Streamlit Page Config is FIRST
st.set_page_config(
    page_title="AI Resume Analyzer", 
    layout="wide", 
    page_icon="📄",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem !important;
        color: #1E88E5;
        text-align: center;
        margin-bottom: 1rem;
    }
    .subheader {
        font-size: 1.5rem;
        color: #0D47A1;
        margin-top: 2rem;
    }
    .card {
        background-color: #f8f9fa;
        border-radius: 10px;
        padding: 20px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin-bottom: 20px;
    }
    .highlight {
        color: #1E88E5;
        font-weight: bold;
    }
    .progress-bar {
        height: 20px;
        border-radius: 10px;
    }
    .skill-tag {
        background-color: #0D47A1;
        color: white;
        padding: 8px;
        margin: 4px;
        border-radius: 5px;
        text-align: center;
        font-weight: 500;
    }
    .template-preview {
        border: 2px solid #1E88E5;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# ✅ Load API Key from .env
load_dotenv()
GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY")

# ✅ Verify API Key
if not GEMINI_API_KEY:
    st.error("🚨 API Key is missing! Ensure GOOGLE_API_KEY is set in the .env file.")
    st.stop()

# ✅ Configure Google Gemini API
genai.configure(api_key=GEMINI_API_KEY)

# ✅ Load NLP Model
@st.cache_resource
def load_nlp_model():
    try:
        return spacy.load("en_core_web_sm")
    except Exception as e:
        st.error("⚠ SpaCy model 'en_core_web_sm' is missing. Run: `python -m spacy download en_core_web_sm`")
        return None

nlp = load_nlp_model()
if not nlp:
    st.stop()

# 📌 Function to Extract Text from PDF
def extract_text_from_pdf(pdf_file):
    try:
        with pdfplumber.open(pdf_file) as pdf:
            text = '\n'.join([page.extract_text() for page in pdf.pages if page.extract_text()])
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

# 📌 Function to Extract Text from DOCX
def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(docx_file)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

# 📌 Function to Parse Resume
def parse_resume(uploaded_file):
    file_extension = uploaded_file.name.split(".")[-1].lower()
    if file_extension == "pdf":
        return extract_text_from_pdf(uploaded_file), file_extension
    elif file_extension == "docx":
        return extract_text_from_docx(uploaded_file), file_extension
    else:
        return None, None

# 📌 Function to Extract Skills using NLP
def extract_skills(text):
    doc = nlp(text.lower())
    skills = set()
    
    # Enhanced common skills list
    common_skills = [
        "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "swift", "kotlin", "go", "rust",
        "sql", "mysql", "postgresql", "mongodb", "oracle", "nosql", "firebase",
        "machine learning", "deep learning", "natural language processing", "nlp", "computer vision",
        "react", "angular", "vue", "node.js", "express", "django", "flask", "spring", "laravel",
        "docker", "kubernetes", "aws", "azure", "gcp", "cloud computing", "devops", "cicd",
        "linux", "unix", "bash", "powershell", "git", "github", "gitlab", "bitbucket",
        "excel", "power bi", "tableau", "data analysis", "data visualization", "statistics", "r",
        "leadership", "communication", "project management", "agile", "scrum", "kanban",
        "html", "css", "sass", "less", "responsive design", "ui/ux", "figma", "sketch",
        "tensorflow", "pytorch", "keras", "pandas", "numpy", "scikit-learn", "matplotlib",
        "rest api", "graphql", "oauth", "authentication", "blockchain", "cybersecurity"
    ]
    
    # Extract single word skills
    for token in doc:
        if token.text in common_skills:
            skills.add(token.text)
    
    # Extract multi-word skills
    for skill in common_skills:
        if len(skill.split()) > 1 and skill.lower() in text.lower():
            skills.add(skill)
    
    return list(skills)

# 📌 Function to Get AI-Powered Resume Suggestions from Gemini API
def ai_resume_improvement_gemini(resume_text, job_description=None):
    try:
        model = genai.GenerativeModel("gemini-1.5-pro")
        
        # Enhanced prompt with job matching if available
        if job_description:
            prompt = f"""
            You are a professional resume consultant. Analyze this resume and provide specific improvements to make it more effective and ATS-friendly. 
            
            RESUME:
            {resume_text}
            
            JOB DESCRIPTION:
            {job_description}
            
            Please provide your analysis in the following format:
            
            ## Overall Assessment
            [Provide a brief overall assessment]
            
            ## Strengths
            - [Strength 1]
            - [Strength 2]
            - [Strength 3]
            
            ## Areas for Improvement
            - [Area 1]
            - [Area 2]
            - [Area 3]
            
            ## Specific Suggestions to Match Job Description
            [Detailed suggestions to better align with the job]
            
            ## Improved Resume
            [Provide a complete, improved version of the resume]
            """
        else:
            prompt = f"""
            You are a professional resume consultant. Analyze this resume and provide specific improvements to make it more effective and ATS-friendly.
            
            RESUME:
            {resume_text}
            
            Please provide your analysis in the following format:
            
            ## Overall Assessment
            [Provide a brief overall assessment]
            
            ## Strengths
            - [Strength 1]
            - [Strength 2]
            - [Strength 3]
            
            ## Areas for Improvement
            - [Area 1]
            - [Area 2]
            - [Area 3]
            
            ## Improved Resume
            [Provide a complete, improved version of the resume]
            """
        
        response = model.generate_content(prompt)
        
        if hasattr(response, 'text'):
            return response.text
        else:
            return "No suggestions available."
    
    except Exception as e:
        return f"Error calling Gemini API: {str(e)}"

# 📌 Function to Generate ATS Score
def calculate_ats_score(resume_text):
    # Initialize base score
    score = 100
    
    # Check for essential sections
    essential_sections = ['experience', 'education', 'skills', 'summary']
    missing_sections = []
    for section in essential_sections:
        if section not in resume_text.lower():
            missing_sections.append(section)
            score -= 10
    
    # Check for proper formatting
    formatting_checks = {
        'bullet_points': len(re.findall(r'•|\u2022|\d+\.', resume_text)) > 5,
        'contact_info': bool(re.search(r'\b[\w\.-]+@[\w\.-]+\.\w{2,4}\b', resume_text)),
        'length': 450 < len(resume_text) < 1500
    }
    
    for check, passed in formatting_checks.items():
        if not passed:
            score -= 5
            
    # Check for keywords
    skills = extract_skills(resume_text)
    if len(skills) < 10:
        score -= (10 - len(skills)//2)
        
    # Ensure score stays within 0-100
    score = max(0, min(100, score))
    
    return {
        "score": score,
        "missing_sections": missing_sections,
        "formatting_issues": [k for k,v in formatting_checks.items() if not v],
        "skills_found": skills
    }

# 📌 Function to extract the improved resume section from AI suggestions
def extract_improved_resume(ai_suggestions):
    if "## Improved Resume" in ai_suggestions:
        parts = ai_suggestions.split("## Improved Resume")
        if len(parts) > 1:
            return parts[1].strip()
    return None

# 📌 Function to convert text to DOCX with template styling
def text_to_docx(text, template_name="Classic"):
    doc = Document()
    
    # Apply template styles
    if template_name == "Classic":
        # Classic template styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Times New Roman'
        heading2_style.font.size = Pt(12)
        heading2_style.font.bold = True
        heading2_style.font.italic = False
    elif template_name == "Modern":
        # Modern template styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(18)
        heading_style.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)  # Blue
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Arial'
        heading2_style.font.size = Pt(14)
        heading2_style.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)  # Darker blue
        heading2_style.paragraph_format.space_after = Pt(6)
    elif template_name == "Professional":
        # Professional template styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # Dark gray-blue
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Calibri'
        heading2_style.font.size = Pt(14)
        heading2_style.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)  # Green
    elif template_name == "Minimalist":
        # Minimalist template styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        style.paragraph_format.line_spacing = 1.2
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = False
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(6)
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Arial'
        heading2_style.font.size = Pt(12)
        heading2_style.font.italic = True
    
    # Process text content
    paragraphs = text.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            if para.strip().startswith('#'):
                level = len(re.match(r'^#+', para.strip()).group(0))
                heading_text = para.strip().lstrip('#').strip()
                doc.add_heading(heading_text, level=level if level <= 9 else 1)
            else:
                lines = para.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith(('- ', '* ', '• ')):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif re.match(r'^\d+\.\s', line):
                        doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
                    else:
                        doc.add_paragraph(line)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ✅ Sidebar for app navigation
with st.sidebar:
    st.image("https://img.icons8.com/color/96/000000/resume.png", width=100)
    st.title("Resume Analyzer")
    st.markdown("---")
    st.markdown("### How it works")
    st.info(
        """
        1. Upload your resume (PDF/DOCX)
        2. Get instant ATS score
        3. Receive AI-powered suggestions
        4. Choose template and download
        """
    )
    st.markdown("---")
    st.markdown("### About")
    st.info(
        """
        AI-powered resume analyzer with:
        - Instant ATS scoring
        - Skills extraction
        - Resume improvement suggestions
        - Professional templates
        """
    )

# ✅ Main App UI
st.markdown("<h1 class='main-header'>📄 AI-Powered Resume Analyzer</h1>", unsafe_allow_html=True)

# Create tabs for different sections
tab1, tab2, tab3 = st.tabs(["📤 Upload & Analyze", "📊 ATS Score", "✨ AI Improvements"])

with tab1:
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    # 📂 Resume Upload
    uploaded_file = st.file_uploader("📂 Upload Your Resume (PDF/DOCX)", type=["pdf", "docx"])
    
    if uploaded_file:
        # Store the original file
        if "original_file" not in st.session_state:
            file_bytes = uploaded_file.getvalue()
            st.session_state.original_file = io.BytesIO(file_bytes)
            st.session_state.original_filename = uploaded_file.name
        
        # Extract Resume Text
        if "resume_text" not in st.session_state or st.session_state.get("resume_filename") != uploaded_file.name:
            with st.spinner("📄 Extracting resume content..."):
                resume_text, file_extension = parse_resume(uploaded_file)
                st.session_state.resume_text = resume_text
                st.session_state.file_extension = file_extension
                st.session_state.resume_filename = uploaded_file.name
                
                # Extract Skills
                st.session_state.skills = extract_skills(resume_text) if resume_text else []
                
                # Reset AI suggestions when a new file is uploaded
                if "ai_suggestions" in st.session_state:
                    del st.session_state.ai_suggestions
                if "improved_resume" in st.session_state:
                    del st.session_state.improved_resume
                
            st.success("✅ Resume uploaded successfully!")
        
        # Display extracted text
        st.markdown("<h2 class='subheader'>📄 Resume Content</h2>", unsafe_allow_html=True)
        st.text_area("Resume Text", st.session_state.resume_text, height=250)
        
        # Display skills
        skills = st.session_state.skills
        if skills:
            st.markdown("<h2 class='subheader'>🔍 Extracted Skills</h2>", unsafe_allow_html=True)
            
            # Create a more visual representation of skills with darker color
            cols = st.columns(3)
            for i, skill in enumerate(skills):
                col_idx = i % 3
                cols[col_idx].markdown(f"<div class='skill-tag'>{skill}</div>", unsafe_allow_html=True)
        
        # AI Analysis button
        if st.button("🚀 Analyze Resume with AI", type="primary", use_container_width=True):
            with st.spinner("🤖 AI is analyzing your resume... This may take a moment..."):
                # Add a small delay to make the spinner visible
                time.sleep(1)
                
                # Get AI suggestions
                ai_suggestions = ai_resume_improvement_gemini(st.session_state.resume_text)
                st.session_state.ai_suggestions = ai_suggestions
                
                # Extract the improved resume part
                improved_resume = extract_improved_resume(ai_suggestions)
                if improved_resume:
                    st.session_state.improved_resume = improved_resume
                
                # Auto-switch to the AI Improvements tab
                st.session_state.active_tab = "AI Improvements"
                st.success("✅ Analysis complete! Check the AI Improvements tab.")
    
    st.markdown("</div>", unsafe_allow_html=True)

with tab2:
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    
    # ATS-specific file uploader
    ats_file = st.file_uploader("📄 Upload Resume for ATS Analysis", 
                               type=["pdf", "docx"],
                               key="ats_upload")
    
    if ats_file:
        with st.spinner("🔍 Analyzing ATS Compatibility..."):
            # Parse resume
            resume_text, _ = parse_resume(ats_file)
            
            if resume_text:
                # Calculate ATS score
                ats_result = calculate_ats_score(resume_text)
                st.session_state.ats_result = ats_result
                
                # Display results
                score = ats_result["score"]
                
                st.markdown("<h2 class='subheader'>📊 ATS Compatibility Score</h2>", unsafe_allow_html=True)
                
                # Visual score meter
                col1, col2 = st.columns([1, 3])
                with col1:
                    st.markdown(f"""
                    <div style="text-align: center;">
                        <div style="font-size: 3rem; font-weight: bold; color: {'#4CAF50' if score >= 70 else '#FFC107' if score >= 50 else '#F44336'};">
                            {score}%
                        </div>
                        <div style="font-size: 0.8rem; color: gray;">
                            ATS Readiness Score
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    st.markdown(f"""
                    <div style="margin-top: 20px;">
                        <div style="background-color: #e0e0e0; border-radius: 10px; height: 20px;">
                            <div style="background-color: {'#4CAF50' if score >= 70 else '#FFC107' if score >= 50 else '#F44336'}; 
                                        width: {score}%; 
                                        height: 100%; 
                                        border-radius: 10px;">
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Score interpretation
                    if score >= 70:
                        st.markdown("✅ **Excellent!** Your resume is ATS-friendly.")
                    elif score >= 50:
                        st.markdown("⚠️ **Moderate.** Some improvements needed for better ATS performance.")
                    else:
                        st.markdown("❌ **Needs Work.** Significant improvements required for ATS systems.")
                
                # Detailed analysis
                st.markdown("#### 🔍 Detailed Analysis")
                
                # Missing sections
                if ats_result["missing_sections"]:
                    st.markdown(f"**Missing Sections:** {', '.join(ats_result['missing_sections'])}")
                
                # Formatting issues
                if ats_result["formatting_issues"]:
                    issues = {
                        'bullet_points': 'Insufficient bullet points',
                        'contact_info': 'Missing contact information',
                        'length': 'Resume length not optimal'
                    }
                    st.markdown("**Formatting Issues:**")
                    for issue in ats_result["formatting_issues"]:
                        st.markdown(f"- {issues.get(issue, issue)}")
                
                # Skills found
                st.markdown("#### ✅ Identified Skills")
                cols = st.columns(3)
                for i, skill in enumerate(ats_result["skills_found"]):
                    cols[i%3].markdown(f"<div class='skill-tag'>{skill}</div>", unsafe_allow_html=True)
                
            else:
                st.error("Error processing uploaded file")

    st.markdown("</div>", unsafe_allow_html=True)

with tab3:
    if uploaded_file and "resume_text" in st.session_state:
        if "ai_suggestions" in st.session_state:
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown("<h2 class='subheader'>✨ AI Improvement Suggestions</h2>", unsafe_allow_html=True)
            
            # Display AI suggestions in a clean format
            st.markdown(st.session_state.ai_suggestions)
            
            # Option to accept changes
            if "improved_resume" in st.session_state and st.session_state.improved_resume:
                st.markdown("<h2 class='subheader'>📝 Accept Improvements</h2>", unsafe_allow_html=True)
                
                # Template Selection
                st.markdown("#### 🎨 Choose a Template")
                selected_template = st.selectbox(
                    "Select a template for your resume:",
                    ["Classic", "Modern", "Professional", "Minimalist"],
                    index=1,
                    key="template_select"
                )
                
                # Template descriptions
                if selected_template == "Classic":
                    st.caption("Traditional format with Times New Roman font, suitable for conservative industries.")
                elif selected_template == "Modern":
                    st.caption("Clean design with Arial font and blue accents, ideal for tech and creative roles.")
                elif selected_template == "Professional":
                    st.caption("Balanced layout with Calibri font, perfect for corporate environments.")
                elif selected_template == "Minimalist":
                    st.caption("Simple and elegant with Arial, focusing on content clarity.")
                
                # Show before/after comparison
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("##### Original Resume")
                    st.text_area("Original", st.session_state.resume_text, height=200, key="original_resume")
                
                with col2:
                    st.markdown("##### Improved Resume")
                    st.text_area("Improved", st.session_state.improved_resume, height=200, key="improved_resume")
                
                # Accept changes and download
                if st.button("✅ Accept Changes & Download DOCX", type="primary", use_container_width=True):
                    # Convert improved resume to DOCX
                    docx_buffer = text_to_docx(st.session_state.improved_resume, selected_template)
                    
                    # Get original filename without extension
                    filename = os.path.splitext(st.session_state.original_filename)[0]
                    current_date = datetime.now().strftime("%Y-%m-%d")
                    new_filename = f"{filename}_improved_{current_date}_{selected_template.lower()}.docx"
                    
                    # Offer for download
                    st.download_button(
                        label="📥 Download Improved Resume (DOCX)",
                        data=docx_buffer,
                        file_name=new_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            # Download original resume button
            if st.session_state.file_extension:
                mime_types = {
                    "pdf": "application/pdf",
                    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                }
                
                mime_type = mime_types.get(st.session_state.file_extension, "application/octet-stream")
                
                st.download_button(
                    label=f"📥 Download Original Resume ({st.session_state.file_extension.upper()})",
                    data=st.session_state.original_file,
                    file_name=st.session_state.original_filename,
                    mime=mime_type
                )
            
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.info("🚀 Click 'Analyze Resume with Gemini AI' in the Upload tab to get improvement suggestions.")
    else:
        st.info("📄 Upload your resume first to get AI-powered improvement suggestions.")

# Footer
st.markdown("---")

st.markdown("<div style='text-align: center; color: gray; font-size: 0.8rem;'>© 2025 AI Resume Analyzer | Powered by Google Gemini</div>", unsafe_allow_html=True)